from docx import Document
from copy import deepcopy
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
import numpy as np
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 配置matplotlib使用中文字体
plt.rcParams['font.sans-serif'] = ['SimHei']  # 设置默认字体为黑体
plt.rcParams['axes.unicode_minus'] = False     # 解决负号显示问题

def delete_paragraphs(doc, text, exact_match=False):
    """
    删除文档中包含指定文本的段落
    :param doc: Document对象
    :param text: 要查找的文本，可以是字符串或字符串列表
    :param exact_match: 是否需要精确匹配。True时只删除完全匹配的段落，False时删除包含文本的段落
    :return: 被删除的段落数量
    """
    if isinstance(text, str):
        text = [text]
    
    paragraphs_to_delete = []
    
    # 遍历所有段落，标记需要删除的段落
    for paragraph in doc.paragraphs:
        for search_text in text:
            if (exact_match and paragraph.text == search_text) or \
               (not exact_match and search_text in paragraph.text):
                paragraphs_to_delete.append(paragraph._element)
                break
    
    # 删除标记的段落
    for paragraph in paragraphs_to_delete:
        paragraph.getparent().remove(paragraph)
    
    # 保存修改后的文档
    doc.save(output_path)
    return len(paragraphs_to_delete)

def replace_template_variables(template_path, output_path, replacements, table_data=None, charts=None):
    """
    替换Word模板中的变量
    :param template_path: 模板文件路径
    :param output_path: 输出文件路径
    :param replacements: 要替换的变量字典
    :param table_data: 表格数据，可以是DataFrame或字典格式{表格索引: DataFrame}
    :param charts: 图表配置字典，格式为{占位符: {'data': DataFrame, 'x_column': str, 'y_column': str, 
                  'title': str, 'xlabel': str, 'ylabel': str}}
    """
    doc = Document(template_path)
    
    remaining_keys = set(replacements.keys())
    
    # 遍历文档中的所有段落
    for paragraph in doc.paragraphs:
        if not remaining_keys:  # 如果所有变量都已替换完成，直接退出
            break
            
        for key in list(remaining_keys):
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, str(replacements[key]))
                remaining_keys.remove(key)
    
    # 如果还有未替换的变量，继续在表格中查找
    if remaining_keys:
        for table in doc.tables:
            if not remaining_keys:  # 如果所有变量都已替换完成，直接退出
                break
                
            for row in table.rows:
                if not remaining_keys:
                    break
                    
                for cell in row.cells:
                    if not remaining_keys:
                        break
                        
                    for paragraph in cell.paragraphs:
                        if not remaining_keys:
                            break
                            
                        for key in list(remaining_keys):
                            if key in paragraph.text:
                                for run in paragraph.runs:
                                    run.text = run.text.replace(key, str(replacements[key]))
                                remaining_keys.remove(key)
                                
    
    # 处理动态表格数据
    if table_data is not None:
        # 如果是单个DataFrame，转换为字典格式
        if isinstance(table_data, pd.DataFrame):
            table_data = {0: table_data}
        
        # 处理每个表格的数据
        for table_index, data in table_data.items():
            if table_index < len(doc.tables):
                table = doc.tables[table_index]
                if len(table.rows) > 1:  # 确保表格至少有一个模板行
                    template_row = table.rows[1]  # 使用第二行作为模板（第一行通常是表头）
                    
                    # 删除原有的模板行
                    for i in range(len(table.rows)-1, 0, -1):
                        table._element.remove(table.rows[i]._element)
                    
                    # 处理DataFrame或列表数据
                    rows_data = data.values.tolist() if isinstance(data, pd.DataFrame) else data
                    
                    # 根据数据添加新行
                    for row_data in rows_data:
                        new_row = deepcopy(template_row._element)
                        table._element.append(new_row)
                        
                        # 获取新添加的行
                        added_row = table.rows[-1]
                        
                        # 填充数据
                        for cell_index, cell_value in enumerate(row_data):
                            if cell_index < len(added_row.cells):
                                cell = added_row.cells[cell_index]
                                for paragraph in cell.paragraphs:
                                    paragraph.text = str(cell_value)
    
    
    insert_charts(doc, charts)
    # 保存修改后的文档
    doc.save(output_path)
    return output_path

def insert_charts(doc, charts=None):
    # 处理图表数据
    if charts is not None:
        for placeholder, config in charts.items():
            # 验证配置
            if not isinstance(config, dict) or 'data' not in config:
                continue
                
            data = config['data']
            if not isinstance(data, pd.DataFrame):
                continue
            
            # 创建图形和坐标轴
            fig, ax = plt.subplots(figsize=(10, 3))

            x_col = config.get('x_column')
            y_col = config.get('y_column')

            if x_col in data.columns and y_col in data.columns:
                # 创建柱状图
                x = np.arange(len(data[x_col]))
                bars = ax.bar(x, data[y_col], width=0.6)
                # 设置x轴刻度和标签
                ax.set_xticks(x)
                ax.set_xticklabels(data[x_col], rotation=45)
                
                # 设置标题和标签
                ax.set_title(config.get('title', ''), fontsize=16, pad=15)
                ax.set_xlabel(config.get('xlabel', x_col), fontsize=12)
                ax.set_ylabel(config.get('ylabel', y_col), fontsize=12)
                
                # 添加数值标签
                for bar in bars:
                    height = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width()/2., height,
                            f'{int(height)}',
                            ha='center', va='bottom')

                # 调整布局
                plt.tight_layout()
                # 添加网格线
                # plt.grid(True, linestyle='--', alpha=0.7)
                ax.grid(axis='y', linestyle='--', alpha=0.7)

                # 将图表保存到内存中
                img_stream = BytesIO()
                plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
                plt.close()
                img_stream.seek(0)
                
                # 在文档中查找并替换图片占位符
                for paragraph in doc.paragraphs:
                    if placeholder in paragraph.text:
                        # 获取当前段落的文本（可能包含其他内容）
                        current_text = paragraph.text.replace(placeholder, "").strip()
                        
                        # 清除当前段落的内容
                        paragraph.clear()

                        # 如果有其他文本，重新添加
                        if current_text:
                            paragraph.add_run(current_text)

                        # 在当前位置后添加新段落
                        new_paragraph = doc.add_paragraph()
                        doc._body._body.insert(doc._body._body.index(paragraph._p) + 1, new_paragraph._p)
                        
                        # 设置新段落对齐方式为居中
                        new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # 添加图片到段落
                        run = new_paragraph.add_run()
                        run.add_picture(img_stream, width=Inches(6))
                        break

    # 示例使用
if __name__ == '__main__':
    template_path = 'templates/monthly_report.docx'
    output_path = 'outputs/monthly_report_filled.docx'
    
    # 示例数据
    replacements = {
        '{$total_a}': '3',
        '{$total_c}': '123',
        '{$total_d}': '456'
    }
    
    # 表格数据示例
    table_data = pd.DataFrame({
        '姓名': ['张三', '李四', '王五'],
        '年龄': [30, 25, 35],
        '职位': ['工程师', '设计师', '经理']
    })
    
    # 替换变量并更新表格
    replace_template_variables(
        template_path=template_path,
        output_path=output_path,
        replacements=replacements,
        table_data=table_data
    )
    
    # 示例1：删除包含特定文本的段落
    delete_count = delete_paragraphs(doc, "亡人事故辖区街道")
    print(f"删除了 {delete_count} 个段落")
    
    # 示例2：删除多个指定文本的段落
    texts_to_delete = ["亡人事故辖区中队分布情况。", "亡人事故辖区街道。", "亡人事故道路分布情况。", "亡人事故时间。"]
    delete_count = delete_paragraphs(doc, texts_to_delete)
    print(f"删除了 {delete_count} 个段落")
    
    # 示例3：精确匹配删除
    delete_count = delete_paragraphs(doc, "完整的段落文本", exact_match=True)
    print(f"精确匹配删除了 {delete_count} 个段落")
    
    # 填充变量示例
    replacements = {
        '{$total_a}': '3',
        '{$total_c}': '123',
        '{$total_d}': '456'
    }
    
    # 表格数据示例
    df = pd.DataFrame({
        '姓名': ['张三', '李四', '王五'],
        '年龄': [30, 25, 35],
        '职位': ['工程师', '设计师', '经理']
    })

    # 准备数据和配置
    df = pd.DataFrame({
        '区域': ['东区', '西区', '南区', '北区'],
        '事故数': [5, 7, 3, 6]
    })

    # 图表配置（数据和配置信息集成在一起）
    charts = {
        '{$accident_chart}': {
            'data': df,  # 直接包含数据
            'x_column': '区域',
            'y_column': '事故数',
            'title': '区域事故统计',
            'xlabel': '区域名称',
            'ylabel': '事故数量'
        }
    }
    
    # 处理变量替换和表格数据
    # replace_template_variables(template_path, output_path, replacements, df)
